
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Cm
import csv
from PIL import Image
from io import BytesIO
from aux_calculations import convert_to_int,conv_vel_a_m,text_mass_hypertrophy,text_diam_LV,text_atrium,remove_signs
import re
import os


#extraer los datos de las tablas
def extract_patient_info(doc)->dict:
    '''
    Accepts word docx and extracts info from the first table where the patient data resides
    returns a dictionary. Enhanced for LibreOffice-converted documents.
    '''
    data = {}
    print(f"[DEBUG] extract_patient_info: Total tables found: {len(doc.tables)}")
    
    # Try multiple tables in case LibreOffice changes table order
    table_indices_to_try = [1, 0, 2] if len(doc.tables) > 2 else [1, 0] if len(doc.tables) > 1 else [0]
    
    for table_idx in table_indices_to_try:
        if table_idx >= len(doc.tables):
            continue
            
        try:
            table = doc.tables[table_idx]
            print(f"[DEBUG] Trying table {table_idx}: {len(table.rows)} rows, {len(table.rows[0].cells) if table.rows else 0} cols")
            
            table_data = {}
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    cell_text = cell.text.strip()
                    if cell_text:
                        print(f"[DEBUG] Table {table_idx}[{i},{j}]: {repr(cell_text)}")
                        
                        # Look for key:value patterns with flexible parsing
                        if ':' in cell_text:
                            try:
                                # Handle multiple colons by taking first split
                                parts = cell_text.split(':', 1)
                                if len(parts) == 2:
                                    key = parts[0].strip().replace(' ', '_').replace('\n', '').replace('\t', '')
                                    value = parts[1].strip().replace('  ', ' ').replace('\n', ' ').replace('\t', ' ')
                                    
                                    # Skip empty values and very short keys
                                    if value and len(key) > 1 and key.lower() not in ['', 'table', 'cell']:
                                        table_data[key] = value
                                        print(f"[DEBUG] Extracted from table {table_idx}: {key} = {value}")
                            except Exception as e:
                                print(f"[DEBUG] Error parsing cell '{cell_text}': {e}")
                                
            # If we found patient data, use this table
            if any(key.lower() in ['name', 'patient_id', 'exam_date', 'gender', 'age'] 
                   for key in table_data.keys()):
                data.update(table_data)
                print(f"[DEBUG] Found patient data in table {table_idx}")
                break
            elif table_data:
                # Keep data from this table but continue searching
                data.update(table_data)
                
        except (IndexError, AttributeError) as e:
            print(f"[DEBUG] Error accessing table {table_idx}: {e}")
            continue
    
    print(f"[DEBUG] Final patient data: {data}")
    return data



def update_dictionary(dic:dict)->dict:
    '''
    takes a dictionary and searches calculations within the values to turn them into new keys
    '''
    calc_list=['*Dimensionless Index','*Flow Rate AS','CSA(LVOT)','SV(LVOT)','CSA(AV SV)','Reg Vol(PISA TR)',
               'EROA(PISA TR)','Flow Rate(PISA TR)','Reg Vol(PISA MR)','EROA(PISA MR)','Flow Rate(PISA MR)',
               'AVA(VTI)','RWT(2D)','LVd Mass(2D-ASE)','MV E/A Ratio',"Average E'","E/Med E'",'LVIDd Index(2D)',
                '%LVPW(2D)','LVESV(A4C Simp)','EF(A4C Simp)',"E/Lat E'","E/Avg E'",'*Aortic Sinus Indexed',
               'LVEDVI(A4C Simp)','LA ESVI(BP A-L)','AVAI(AVA VTI)','SI(LVOT)','%IVS(2D)'
              ]
    updates={}
    for val in dic.values():
        for c in calc_list:
            if c in val:
                index = val.index(c)
                if index + 1 < len(val):  # Check bounds
                    key=val[index]
                    value=val[index+1]
                    updates[key]=value
    dic.update(updates)
    return dic

def dic_cleaning(data)->dict:
    '''
    gets a dict. All the values are set to string, if there are more than one value gets the first.
    if there are keys within a value list it removes that and the following number that is the value of that key
    '''

    for key, value in data.items():
        if isinstance(value, list):
            if len(value) == 1:
                data[key] = value[0]
            else:
                key_in_value = [v for v in value if v in data]
                if key_in_value:
                    index = value.index(key_in_value[0])
                    data[key] = value[:index][0] if index > 0 else ''
    return data

def get_measure_table(doc)->'docx.table.Table | None':
    '''
    Accepts word docx and extracts the table object where the measurements are
    '''
    print("[DEBUG] get_measure_table: Searching for measurements table...")
    for i, table in enumerate(doc.tables):
        print(f"[DEBUG] Table {i}: {len(table.rows)} rows")
        for j, row in enumerate(table.rows[:2]):  # Check first 2 rows
            for k, cell in enumerate(row.cells):
                if cell.text.strip():
                    print(f"[DEBUG]   [{j},{k}]: {repr(cell.text)}")
        # More flexible search for measure table
        if any('measure' in cell.text.lower() for row in table.rows for cell in row.cells):
            print(f"[DEBUG] Found measurements table at index {i}")
            return table
    print("[DEBUG] No measurements table found")
    return None

def get_mot_table(doc)->'docx.table.Table | None':
    print("[DEBUG] get_mot_table: Searching for WMS table...")
    for i, table in enumerate(doc.tables):
        # More flexible search for WMS table
        if any('wms' in cell.text.lower() for row in table.rows for cell in row.cells):
            print(f"[DEBUG] Found WMS table at index {i}")
            return table
    print("[DEBUG] No WMS table found")
    return None

def mot_extractor(table)->dict:

    """
    Extracts wall motion scores from a nested table.

    Returns:
        dict: With structure {'mot': [{'key': segment_name, 'motilidad': [rest, peak, recovery]}]}
    """

    mot={}
    for index_r, row in enumerate(table.rows):
        for cell in row.cells:
             for inner_table in cell.tables:
                    for index_r,inner_row in enumerate(inner_table.rows):
                       #la fila 1 contiene la primera info,la ultima info en la 17 (apex) 
                        if 1 <= index_r <= 17:
                            key = None
                            values = []
                            for index_c, inner_cell in enumerate(inner_row.cells):
                                #las primeras cuatro celdas son segment ID
                                #la 5 celda es el nombre del segmento
                                #la 6 es baseline, 7 peak, 8 recovery 
                                if index_c == 5:
                                    key = inner_cell.text.lower()
                                if 5 < index_c <= 8 and key:
                                    try:
                                        values.append(int(inner_cell.text))
                                    except ValueError as e:
                                        print(f'{e} no se pudo convertir motilidad en segmento {key}')
                                        values.append(inner_cell.text)
                            if key and values:  # Store the key-value pair only if both key and values exist
                                mot[key] = values
                                            

    return {'mot': [{'key': k, 'motilidad': v} for k, v in mot.items()]} 

def get_measurements(table,gender):
    data={}
    units=['mm','cm','ml','g','ms','mmHg','cm²','cm/s','ml/s','cm²','m²','ml/m²','cm²/m²','g/m²']
    
    print(f"[DEBUG] get_measurements: Processing table with {len(table.rows)} rows")
    
    # Check if this is a LibreOffice flattened table (no nested tables)
    has_nested_tables = any(len(cell.tables) > 0 for row in table.rows for cell in row.cells)
    print(f"[DEBUG] Table has nested tables: {has_nested_tables}")
    
    if not has_nested_tables:
        # Handle LibreOffice flattened structure
        print("[DEBUG] Using LibreOffice flat table parsing")
        return parse_flattened_measurements(table, gender)
    else:
        # Use original nested table parsing
        print("[DEBUG] Using original nested table parsing")
        return parse_nested_measurements(table, gender)

def parse_flattened_measurements(table, gender):
    """Parse measurements from LibreOffice-flattened table structure with enhanced detection"""
    data = {}
    print(f"[DEBUG] Parsing flattened table: {len(table.rows)} rows")
    
    # Enhanced field mapping with more variations
    field_mapping = {
        'diámetro diastólico del vi': 'LVIDd',
        'diámetro sistólico del vi': 'LVIDs', 
        'espesor diastólico del septum': 'IVSd',
        'espesor diastólico de la pared': 'LVPWd',
        'masa vi': 'LV_Mass',
        'raíz de aorta': 'Aortic_Root',
        'aurícula izquierda': 'LA_Area',
        'aurícula derecha': 'RA_Area',
        'diámetro basal vd': 'RV_Basal',
        'fac%': 'FAC',
        'tsvi': 'TSVI',
        # Add velocity terms for second table
        'vel pico': 'Vel_Peak',
        'grad pico': 'Grad_Peak',
        'aortica': 'Aortic_Vel',
        'mitral': 'Mitral_Vel',
        'tricúspide': 'Tricuspid_Vel',
        'pulmonar': 'Pulmonary_Vel'
    }
    
    # First pass: create a cell content map for cross-referencing
    cell_map = {}
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:
                cell_map[(row_idx, cell_idx)] = cell_text
                print(f"[DEBUG] Cell map [{row_idx},{cell_idx}]: {repr(cell_text)}")
    
    # Second pass: enhanced value search
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip().lower()
            
            if cell_text:
                # Look for medical field names
                for pattern, key in field_mapping.items():
                    if pattern in cell_text:
                        value = find_associated_value(cell_map, row_idx, cell_idx, len(row.cells), len(table.rows))
                        
                        if value:
                            unit = find_associated_unit(cell_map, row_idx, cell_idx, len(row.cells), len(table.rows))
                            data[key] = {'value': value, 'unit': unit or ''}
                            print(f"[DEBUG] Enhanced extraction: {key} = {value} {unit or ''}")
                        break
    
    # Third pass: look for standalone numeric values that might be measurements
    standalone_values = find_standalone_numeric_values(cell_map, len(table.rows))
    if standalone_values:
        print(f"[DEBUG] Found {len(standalone_values)} standalone numeric values")
        for pos, value in standalone_values.items():
            print(f"[DEBUG] Standalone value at {pos}: {value}")
    
    print(f"[DEBUG] Enhanced flattened parsing extracted: {data}")
    return data

def find_associated_value(cell_map, row_idx, cell_idx, max_cols, max_rows):
    """Find numeric value associated with a medical field using multiple search strategies"""
    
    # Strategy 1: Check immediate right cell
    if (row_idx, cell_idx + 1) in cell_map:
        right_text = cell_map[(row_idx, cell_idx + 1)]
        if is_numeric_value(right_text):
            return extract_numeric_value(right_text)
    
    # Strategy 2: Check same row, further right cells
    for check_col in range(cell_idx + 2, min(cell_idx + 4, max_cols)):
        if (row_idx, check_col) in cell_map:
            check_text = cell_map[(row_idx, check_col)]
            if is_numeric_value(check_text):
                return extract_numeric_value(check_text)
    
    # Strategy 3: Check next row, same column and nearby
    for check_row in range(row_idx + 1, min(row_idx + 3, max_rows)):
        for check_col in range(max(0, cell_idx - 1), min(cell_idx + 3, max_cols)):
            if (check_row, check_col) in cell_map:
                check_text = cell_map[(check_row, check_col)]
                if is_numeric_value(check_text):
                    return extract_numeric_value(check_text)
    
    return None

def find_associated_unit(cell_map, row_idx, cell_idx, max_cols, max_rows):
    """Find unit associated with a medical field"""
    units = ['mm', 'cm', 'ml', 'g', 'mmHg', 'm²', 'cm²', 'cm/s', 'ml/s', 'ml/m²', 'cm²/m²', 'g/m²', 'ms', '%']
    
    # Search in nearby cells for units
    for check_row in range(max(0, row_idx - 1), min(row_idx + 3, max_rows)):
        for check_col in range(max(0, cell_idx - 1), min(cell_idx + 4, max_cols)):
            if (check_row, check_col) in cell_map:
                check_text = cell_map[(check_row, check_col)]
                for unit in units:
                    if unit in check_text:
                        return unit
    return None

def find_standalone_numeric_values(cell_map, max_rows):
    """Find all numeric values in the table that might be measurements"""
    standalone_values = {}
    
    for (row_idx, cell_idx), text in cell_map.items():
        if is_numeric_value(text) and not any(char.isalpha() for char in text if char not in '.,'):
            # This is a pure numeric value
            numeric_val = extract_numeric_value(text)
            if numeric_val and float(numeric_val) > 0:  # Positive meaningful values
                standalone_values[(row_idx, cell_idx)] = numeric_val
    
    return standalone_values

def extract_numeric_value(text):
    """Extract the actual numeric value from text"""
    if not text:
        return None
    # Extract numbers with decimals
    import re
    match = re.search(r'\d+(?:[.,]\d+)?', text.replace(',', '.'))
    if match:
        return match.group()
    return None

def parse_nested_measurements(table, gender):
    """Original nested table parsing (for non-LibreOffice documents)"""
    data={}
    units=['mm','cm','ml','g','ms','mmHg','cm²','cm/s','ml/s','cm²','m²','ml/m²','cm²/m²','g/m²']
    
    for row_idx, row in enumerate(table.rows):
        print(f"[DEBUG] Processing row {row_idx}: {len(row.cells)} cells")
        for cell_idx, cell in enumerate(row.cells):
            print(f"[DEBUG] Cell [{row_idx},{cell_idx}]: {len(cell.tables)} nested tables, text: {repr(cell.text[:100])}")
            for st_idx, st in enumerate(cell.tables):
                for index_rs,rs in enumerate(st.rows):
                    if index_rs>1:
                        elements=[]
                        values=[]
                        for index_cs,cs in enumerate(rs.cells):
                            if index_cs==0:
                                text=cs.text.replace(' ','',2)
                                if not text.startswith(' '):
                                    key=text
                                    subkey=''
                                    elements.append(key)
                                else:
                                    subkey=text
                                    elements.append(subkey.strip())   
                            if (cs.text.strip() 
                                not in elements 
                                and (cs.text.strip()!='') 
                                and not (cs.text.endswith('Last'))
                                and cs.text.strip() not in units
                               ):
                                value=cs.text.strip()
                                values.append(value)
                                data[key+subkey]=values
    
    data=update_dictionary(data) 
    data=dic_cleaning(data)
    data=convert_to_int(data)
    data=conv_vel_a_m(data)
    data['Gender']=gender
    data=text_mass_hypertrophy(data)
    data=text_diam_LV(data)
    data=text_atrium(data)
    data=remove_signs(data) #tiene que ir ultimo porque las funciones de interpretaci'on usan el dict como se ve en el word del equpo.
    return data




def image_extractor(doc, template, tipo, image_width=Cm(8), image_height=Cm(5.36)) -> dict:
    '''
    Extrae las imágenes del reporte docx del dispositivo Vinno y devuelve un diccionario con objetos InlineImage.
    Requiere:
        - doc: Documento a extraer
        - template: Template donde se renderizarán las imágenes
        - tipo: Tipo de template (por ejemplo, 'stress')
    Se establecen medidas habituales de 5.36x8 cm, excepto para el mapa polar del stress que es 8.22 x 16.23 y 6.39 x 16.23 cm.
    '''
    image_dict = {}

    # Extraer imágenes
    for rel in doc.part.rels:
        rel_obj = doc.part.rels[rel]
        if 'image' in rel_obj.reltype:
            image_data = rel_obj.target_part.blob
            image = Image.open(BytesIO(image_data))
            compressed_image = BytesIO()

            # Extraer el número de imagen
            target = rel_obj.target_ref.split('.')[0].replace(r'media/', '')
            image_number = int(target.replace('image', ''))

            # Si es 'stress' y la imagen es 1 o 2, se manejan de manera diferente
            if tipo == 'stress' and image_number in [1, 2]:
                compressed_image = BytesIO(image_data)  # Mantener PNG sin convertir
                compressed_image.seek(0)
                # Definir tamaños específicos para las primeras dos imágenes
                if image_number == 1:
                    image_dict[target] = InlineImage(template,
                                                     compressed_image,
                                                     width=Cm(16.23), height=Cm(8.22))
                else:
                    image_dict[target] = InlineImage(template,
                                                     compressed_image,
                                                     width=Cm(16.23), height=Cm(6.39))
            else:
                # Si no es 'stress' o es una imagen normal, convertir a JPEG si es necesario
                if image.format != 'JPEG':
                    image = image.convert("RGB")  # Convertir a RGB para JPEG
                image.save(compressed_image, format='JPEG', quality=85)  # Guardar como JPEG
                compressed_image.seek(0)

                # Asignar tamaño predeterminado
                image_dict[target] = InlineImage(template,
                                                 compressed_image,
                                                 width=image_width,
                                                 height=image_height)

    # Ordenar las imágenes por su número
    key = sorted(image_dict.keys(), key=lambda image_name: int(image_name.replace('image', '')))
    image_dict = {i: image_dict[i] for i in key}
    image_dict = {'image': [{'key': k, 'image': v} for k, v in image_dict.items()]}

    return image_dict

def is_numeric_value(text):
    """Check if text contains a numeric value"""
    if not text:
        return False
    # Remove common non-numeric characters and check if what remains is numeric
    cleaned = re.sub(r'[^\d.,]', '', text.strip())
    if not cleaned:
        return False
    try:
        float(cleaned.replace(',', '.'))
        return True
    except ValueError:
        return False

def extract_unit(text):
    """Extract unit from text"""
    units = ['mm', 'cm', 'ml', 'g', 'mmHg', 'm²', 'cm²', 'cm/s', 'ml/s', 'ml/m²', 'cm²/m²', 'g/m²', 'ms']
    for unit in units:
        if unit in text:
            return unit
    return None

#####

def mot_grouper(dic,idx): 
    group={
    'Normoquinesia':[],
    'Hipoquinesia':[],
    'Aquinesia':[],
    'Disquinesia':[],
    'Aneurismático':[],
    }        
    for dic in dic['mot']:
        if dic['motilidad'][idx]==2:
            group['Hipoquinesia'].append(dic['key'])
        elif dic['motilidad'][idx]==3:
            group['Aquinesia'].append(dic['key'])    
        elif dic['motilidad'][idx]==4:
            group['Disquinesia'].append(dic['key'])  
        elif dic['motilidad'][idx]==5:
            group['Aneurismático'].append(dic['key'])
        else:
            group['Normoquinesia'].append(dic['key'])

    return group

def mot_interpreter(dic:dict)->dict:
    """
    groups myocardial segments by motility

    """
    segments_group={}
    segments_group['reposo']=mot_grouper(dic,0)
    segments_group['esfuerzo']=mot_grouper(dic,1)
    return segments_group


def delta_motility(dic: dict)-> dict:

    improve=[]
    ischaemic=[]
    unchanged=[]
    for i in dic['mot']:
        segment=i['key']
        score_rep=i['motilidad'][0]
        score_stress=i['motilidad'][1]
        delta=score_rep-score_stress
        if delta <0:
            ischaemic.append(segment)
        elif delta >0:
            improve.append(segment)
        else: unchanged.append(segment)
    delta_mot={'improve':improve,
              'ischaemic':ischaemic,
              'unchanged':unchanged}
    return delta_mot






# In[ ]:


def process_pdf_images(image_paths, template, tipo, image_width=Cm(8), image_height=Cm(5.36)) -> dict:
    """
    Processes images extracted from PDF for use with DocxTemplate.

    Args:
        image_paths: Dictionary of image paths from PDF extraction
        template: DocxTemplate instance
        tipo: Template type
        image_width: Default image width
        image_height: Default image height

    Returns:
        Dictionary with InlineImage objects
    """
    image_dict = {}

    for name, path in image_paths.items():
        if os.path.exists(path):
            with open(path, 'rb') as f:
                image_data = f.read()

            image = Image.open(BytesIO(image_data))
            compressed_image = BytesIO()

            # Extract image number from name if possible
            image_number = 0
            match = re.search(r'img_(\d+)', name)
            if match:
                image_number = int(match.group(1))

            # Handle stress template special cases
            if tipo == 'stress' and image_number in [1, 2]:
                compressed_image = BytesIO(image_data)
                compressed_image.seek(0)
                if image_number == 1:
                    image_dict[f"image{image_number}"] = InlineImage(template,
                                                     compressed_image,
                                                     width=Cm(16.23), height=Cm(8.22))
                else:
                    image_dict[f"image{image_number}"] = InlineImage(template,
                                                     compressed_image,
                                                     width=Cm(16.23), height=Cm(6.39))
            else:
                # Convert to JPEG if needed
                if image.format != 'JPEG':
                    image = image.convert("RGB")
                image.save(compressed_image, format='JPEG', quality=85)
                compressed_image.seek(0)

                image_dict[f"image{image_number if image_number else len(image_dict)+1}"] = InlineImage(
                    template, compressed_image, width=image_width, height=image_height)

    # Format for template
    image_dict = {'image': [{'key': k, 'image': v} for k, v in image_dict.items()]}
    return image_dict

def generate_motility_report(mot):
    """
    Generates a motility report based on mot.

    Args:
        mot: dict returned by the mot extractor

    Returns:
        str: The generated motility report.
    """

    reposo = []
    esfuerzo=[]
    mejoria=[]
    mot_interpret=mot_interpreter(mot)
    delta_mot=delta_motility(mot)

    #group motility during reposo
    if len(mot_interpret['reposo'].get("Normoquinesia", [])) != 17:
        for key, segments in mot_interpret['reposo'].items():
            if key != "Normoquinesia":
                new_segments = [segment for segment in segments]
                if new_segments:
                    segments_str = ", ".join(new_segments)
                    reposo.append(f"{key}: {segments_str}.")

    else: reposo.append('Sin trastornos de la motilidad basal.')

    # Group new conditions in esfuerzo compared to reposo
    for key, segments in mot_interpret['esfuerzo'].items():
        if key != "Normoquinesia":
            new_segments = [segment for segment in segments if segment not in mot_interpret['reposo'].get(key, [])]
            if new_segments:
                segments_str = ", ".join(new_segments)
                esfuerzo.append(f"Nueva {key}: {segments_str}.")

#     # Handle segments with delta < 0 (ischaemic)
#     if delta_mot.get("ischaemic"):
#         segments = ", ".join(delta_mot["ischaemic"])
#         report.append(f"Nueva isquemia de los segmentos: {segments}.")

    # Handle segments with delta > 0 (improve)
    if delta_mot.get("improve"):
        segments = ", ".join(delta_mot["improve"])
        mejoria.append(f"Mejoria de los segmentos: {segments}.")

    # Handle unchanged segments
    unchanged = delta_mot.get("unchanged", [])
    if len(unchanged) == 17 and len(mot_interpret['reposo'].get("Normoquinesia", [])) == 17:
        esfuerzo.append("Hipercontractilidad de todos los segmentos.")
    elif len(unchanged) == 17 and len(mot_interpret['reposo'].get("Normoquinesia", [])) != 17:
        esfuerzo.append("Sin nuevos trastornos de motilidad.")
    reposo=' '.join(reposo)
    esfuerzo=' '.join(esfuerzo)
    mejoria=' '.join(mejoria)

    report={'reposo':reposo,
           'esfuerzo':esfuerzo,
           'mejoria':mejoria}

    return report



