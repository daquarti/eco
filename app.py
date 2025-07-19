from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import openai
from dotenv import load_dotenv
import io
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import os
from PIL import Image
import openai
from dotenv import load_dotenv

# Cargar variables de entorno desde .env para desarrollo local
if os.path.exists('.env'):
    load_dotenv()

app = Flask(__name__)
CORS(app)

# Configurar OpenAI
api_key = os.getenv('OPENAI_API_KEY')
if not api_key:
    raise ValueError('OPENAI_API_KEY environment variable is not set')

openai.api_key = api_key



@app.route('/api/hello', methods=['POST'])
def hello():
    try:
        # Crear un nuevo documento
        doc = Document()
        
        # Agregar un título
        heading = doc.add_heading('Informe de Cardiología', 0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Agregar fecha
        fecha_actual = datetime.now().strftime('%d/%m/%Y')
        p = doc.add_paragraph()
        p.add_run(f'Fecha: {fecha_actual}').bold = True
        doc.add_paragraph()
        
        # Agregar una tabla con mejor formato
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.allow_autofit = True
        
        # Dar formato al encabezado
        header_cells = table.rows[0].cells
        for cell in header_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].add_run()
            run.font.bold = True
            run.font.size = Pt(11)
        
        header_cells[0].text = 'Estudio'
        header_cells[1].text = 'Resultado'
        
        # Agregar contenido de la tabla
        row_cells = table.add_row().cells
        for cell in row_cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[0].text = 'Ecocardiograma'
        row_cells[1].text = 'Normal'
        
        # Guardar el documento en memoria
        f = io.BytesIO()
        doc.save(f)
        f.seek(0)
        
        return send_file(
            f,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name='informe_cardiologia.docx'
        )
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/analyze', methods=['POST'])
def analyze_text():
    try:
        data = request.json
        text = data.get('text', '')
        
        # Llamada a OpenAI
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "Eres un cardiólogo experto. Analiza el siguiente texto médico y proporciona un análisis estructurado que incluya:\n1. Síntomas principales\n2. Posibles diagnósticos diferenciales ordenados por probabilidad\n3. Nivel de urgencia (Bajo/Medio/Alto)\n4. Recomendaciones inmediatas\n\nMantén el análisis conciso pero informativo."},
                {"role": "user", "content": text}
            ]
        )
        
        analysis = response.choices[0].message.content
        return jsonify({"analysis": analysis})
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    app.run(port=5001, debug=True)
