#!/usr/bin/env python
# coding: utf-8

# In[1]:


import os
from docxtpl import DocxTemplate
from docx import Document

def template_selector(path)->tuple: 

    '''
    Carga por default el template de doppler cardiaco en el pendrive
    acepta el path del estudio del paciente generado por el equipo vinno
    Si el path contiene las palabras carotid, arteries o veins cambia al tipo de template correspondiente
    si no encuentra estas palabras abre el docx y busca la palabra WMS en las tablas del informe (esto determinaria un ecostres)
    si la encuentra cambia al ecostress
    Devuelve el template correspondiente con el tipo de template, para ser usado en la funcion render template
    '''
    #por default elijo la plantilla de cardio
    # Usar rutas relativas para compatibilidad cloud/container
    base_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(base_dir, "auto card.docx")
    tipo = 'card'
    path_str = str(path)
    try:
        if 'Carotid' in path_str:
            template_path = os.path.join(base_dir, "auto vc.docx")
            tipo = 'carotid'
        elif 'Arteries' in path_str:
            template_path = os.path.join(base_dir, "auto art.docx")
            tipo = 'art'
        elif 'Veins' in path_str:
            template_path = os.path.join(base_dir, "auto ven.docx")
            tipo = 'ven'
        else:
            doc = Document(path)
            if doc.tables[2].rows[0].cells[0].text == 'WMS':
                template_path = os.path.join(base_dir, "auto stress.docx")
                tipo = 'stress'
    except (IndexError, AttributeError) as e:
        print(f"Error: {e}, defaulting to 'card' template.")

    template = DocxTemplate(template_path)
    return template, tipo





def template_selector_gui(path_template,path_study)->tuple: 

    '''
    Takes the path of the templates folder and the path to the individual study, by defaults loads 'auto card'.
    will change to 'auto vc', 'auto art','auto stress' or 'auto ven' depending on the extention of the path of the study
    returns the corresponding template objetct and its type
    Devuelve el template correspondiente con el tipo de template, para ser usado en la funcion render template
    '''
    #por default elijo la plantilla de cardio
    template_path=path_template+'\\auto card.docx'
    tipo='card'
    path_str=str(path_study)
    #chqueo si en el path esta la palabra carotida si esta cambio a esa plantilla
    try:
        if 'Carotid' in path_str:

            template_path=path_template+'\\auto vc.docx'  
            tipo='carotid'
        elif 'Arteries' in path_str :#art
            template_path=path_template+'\\auto art.docx'
            tipo='art'
        elif'Veins' in path_str:
            template_path=path_template+'\\auto ven.docx'
            tipo='ven'
        else: 
            #el try es porque para distinguir cardio de stress neceisto acceder a una tabla, si la tabla no esta para el progrmaa
            #de esta manera con un try evito que se rompa si no esta y le asigo el valor default
            doc=Document(path_study)
            if doc.tables[2].rows[0].cells[0].text=='WMS':
                template_path=path_template+'\\auto stress.docx'
                tipo='stress'
    except (IndexError, AttributeError) as e:
        # Log or handle specific exceptions if needed
        print(f"Error: {e}, defaulting to 'card' template.")

    template = DocxTemplate(template_path)
    return template,tipo




# 
# def remove_empty_paragraphs_in_section_stress(template):
#     """
#     Removes empty paragraphs or those with only whitespace between the sections 'Reposo' and 'Conclusión' stress tremplate.
# 
#     Args:
#         template (DocxTemplate): A rendered DocxTemplate object.
# 
#     Modifies:
#         The template object in place.
#     """
#     doc = template.docx  # Access the underlying Document object
#     paragraphs = doc.paragraphs
# 
#     # Flags to track when to start and stop removing paragraphs
#     in_target_section = False
# 
#     for paragraph in paragraphs:
#         # Check for section start
#         if "Reposo" in paragraph.text:
#             in_target_section = True
#         # Check for section end
#         elif "Conclusión" in paragraph.text:
#             in_target_section = False
# 
#         # Remove empty paragraphs only within the target section
#         if in_target_section and not paragraph.text.strip():
#             p_element = paragraph._element
#             p_element.getparent().remove(p_element)



